import docx
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.section import WD_ORIENTATION
import os
from pdf2docx import Converter
import re
from copy import deepcopy
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from PIL import Image
import fitz  # PyMuPDF

def convert_pdf_to_docx(pdf_path, output_docx=None):
    """
    Convert PDF to DOCX format using pdf2docx library.
    
    Args:
        pdf_path (str): Path to the PDF file
        output_docx (str): Optional output path for the DOCX file
        
    Returns:
        str: Path to the converted DOCX file
    """
    try:
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"File not found: {pdf_path}")
        
        if not pdf_path.lower().endswith('.pdf'):
            raise ValueError("File must be a PDF")
        
        # If no output path specified, create one based on input path
        if output_docx is None:
            output_docx = os.path.splitext(pdf_path)[0] + '_converted.docx'
        
        # Convert PDF to DOCX
        cv = Converter(pdf_path)
        cv.convert(output_docx)
        cv.close()
        
        print(f"Successfully converted PDF to DOCX: {output_docx}")
        return output_docx
    
    except Exception as e:
        print(f"Error converting PDF to DOCX: {str(e)}")
        return None

def extract_text_from_docx(file_path):
    """
    Extract text content from a Word document.
    
    Args:
        file_path (str): Path to the Word document
        
    Returns:
        str: Extracted text content
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.lower().endswith(('.docx', '.doc')):
            raise ValueError("File must be a Word document (.doc or .docx)")
        
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                full_text.append(para.text)
        
        return '\n'.join(full_text)
    
    except Exception as e:
        print(f"Error extracting text from document: {str(e)}")
        return None

def parse_document(file_path):
    """
    Parse an existing document and return its text content.
    
    Args:
        file_path (str): Path to the document file
    
    Returns:
        str: The text content of the document
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file extension to determine parsing method
        file_ext = os.path.splitext(file_path.lower())[1]
        
        if file_ext == '.pdf':
            # Convert PDF to DOCX first
            docx_path = convert_pdf_to_docx(file_path)
            if docx_path and os.path.exists(docx_path):
                return extract_text_from_docx(docx_path)
            return None
        elif file_ext in ('.docx', '.doc'):
            # Parse Word document
            return extract_text_from_docx(file_path)
        else:
            raise ValueError("File must be a Word document (.doc or .docx) or a PDF (.pdf)")
    
    except Exception as e:
        print(f"Error parsing document: {str(e)}")
        return None

def set_page_size_and_margins(doc):
    """
    Set the page size to A4 and appropriate margins.
    
    Args:
        doc: The document to modify
    """
    for section in doc.sections:
        # Set page size to A4 (210mm × 297mm)
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.orientation = WD_ORIENTATION.PORTRAIT
        
        # Set margins (2.54cm/1 inch on all sides)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Set header and footer distances
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

def is_heading(text):
    """
    Determine if the text is likely a heading using improved detection rules.
    
    Args:
        text (str): The text to check
        
    Returns:
        bool: True if the text is likely a heading, False otherwise
    """
    if not text.strip():
        return False
    
    # Check text length (headings are typically shorter)
    if len(text.strip()) > 150:  # Increased length threshold
        return False
    
    # New condition: Check if text contains a colon
    if ':' in text:
        # Additional check to ensure it's not a sentence with a colon
        # If the text after the colon is too long, it's probably not a heading
        parts = text.split(':', 1)
        if len(parts) > 1 and len(parts[1].strip()) < 100:
            return True
    
    # New condition: Check if text has less than 6 words
    word_count = len(text.strip().split())
    if word_count < 6:
        return True
    
    # Common heading words at start (case-insensitive)
    heading_words = [
        'introduction', 'conclusion', 'summary', 'overview', 'background',
        'methodology', 'results', 'discussion', 'recommendations', 'references',
        'abstract', 'executive summary', 'scope', 'objectives', 'findings',
        'analysis', 'recommendation', 'appendix', 'bibliography', 'acknowledgments'
    ]
    
    # Check if text starts with any heading word
    text_lower = text.strip().lower()
    if any(text_lower.startswith(word) for word in heading_words):
            return True
    
    # Check for common heading patterns
    heading_patterns = [
        r'^\s*\d+(\.\d+)*\s+',  # Numbered headings (e.g., "1.", "1.1", "1.1.1")
        r'^\s*[IVX]+\.\s+',     # Roman numeral headings (e.g., "I.", "II.", "III.")
        r'^\s*[A-Z]\.\s+',      # Letter headings (e.g., "A.", "B.", "C.")
        r'^[A-Z\s]{3,}$',       # All caps text (common for headings)
        r'^\s*[•\-*]\s+',      # Text with bullet points
        r'^(Section|Chapter|Part|Appendix)\s+\d+',  # Text with specific heading keywords
        r'^[\w\s]{3,30}\.$',   # Text that's very short (3-5 words) and ends with period
        r'^[A-Z\s]{3,30}$',    # Text that's all uppercase and short
        r'^(Overview|Summary|Introduction|Conclusion|Background|Methodology|Results|Discussion|Recommendations|References)',
        r'^\d+\.\s+[A-Za-z]',  # Text with numbers followed by text
        r'(Overview|Summary|Introduction|Conclusion|Background|Methodology|Results|Discussion|Recommendations|References)$'
    ]
    
    # Check if text matches any heading pattern
    return any(re.match(pattern, text.strip(), re.IGNORECASE) for pattern in heading_patterns)

def is_subheading(text, previous_texts=None, heading_detected=False):
    """
    Determine if the text is likely a subheading.
    
    Args:
        text (str): The text to check
        previous_texts (list): List of previous paragraph texts to check context
        heading_detected (bool): Whether a main heading was recently detected
        
    Returns:
        bool: True if the text is likely a subheading, False otherwise
    """
    if not text.strip() or len(text.strip()) > 100:
        return False
    
    # If a heading was recently detected, this could be a subheading
    if heading_detected:
        # Subheadings often start with numbering or bullets at a deeper level
        if re.match(r'^\s*(\d+\.\d+|\d+\.\d+\.\d+|[a-z]\.|\([a-z]\)|\([ivx]+\))', text.strip()):
            return True
        
        # Subheadings may have similar formatting but are typically shorter
        if is_heading(text) and len(text.strip()) < 60:
            return True
    
    # Check for typical subheading patterns regardless of previous heading
    if re.match(r'^\s+\d+\.|\s+[a-z]\.|\s+•|\s+-|\s+\*', text):  # Indented numbering or bullets
        return True
        
    # Check for common subheading prefixes
    common_prefixes = ['subsection', 'part', 'item', 'sub', 'section']
    return any(text.lower().startswith(prefix) and len(text) < 60 for prefix in common_prefixes)

def format_paragraph(paragraph, is_heading_text=False, is_subheading_text=False):
    """
    Apply formatting to a paragraph based on whether it's a heading or normal text.
    
    Args:
        paragraph: The paragraph to format
        is_heading_text: Whether the paragraph is a heading
        is_subheading_text: Whether the paragraph is a subheading
        
    Returns:
        The formatted paragraph
    """
    if is_heading_text:
        # Heading formatting: left alignment, 14pt font size, bold
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(12)  # Standard space after headings
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        # Regular paragraph formatting: left alignment, 12.5pt font size
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(12)
        for run in paragraph.runs:
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    return paragraph

def add_space_after_paragraph(paragraph, is_heading=False, is_subheading=False):
    """
    Adds proper spacing after a paragraph according to Word standards.
    Also applies pagination controls to prevent orphaned headings.
    
    Args:
        paragraph: The paragraph to modify
        is_heading: Whether the paragraph is a heading
        is_subheading: Whether the paragraph is a subheading
        
    Returns:
        The modified paragraph
    """
    # Use Word's standard spacing
    spacing = Pt(18) if is_heading else Pt(14) if is_subheading else Pt(12)
    paragraph.paragraph_format.space_after = spacing
    
        # Keep heading with next paragraph to prevent orphaned headings
    if is_heading or is_subheading:
        paragraph.paragraph_format.keep_with_next = True
    
    return paragraph

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    
    Args:
        cell: The cell to modify
        **kwargs: Border parameters (top, bottom, start, end, etc.)
            Each parameter is a dict with keys: sz (size), val (line style), color, space, shadow
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # List over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # Check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # Set attributes for the border element
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def enhance_table_appearance(table):
    """
    Enhance the appearance of a table with better formatting.
    
    Args:
        table: The table to enhance
        
    Returns:
        The enhanced table
    """
    # Set table properties for better appearance
    try:
        # Set table alignment to center
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set consistent cell padding
        for row in table.rows:
            for cell in row.cells:
                # Set cell margins/padding (in points)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)
                    # Center align text in cells
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Set font properties for text in cells
                    for run in paragraph.runs:
                        run.font.size = Pt(11)  # Slightly smaller font for tables
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Format header row (first row) if table has at least one row
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # Make header text bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        
                    # Add light gray background to header cells
                    # Note: This requires XML manipulation which is complex
                    # We'll use the border function to make headers stand out instead
        
        print("Enhanced table appearance with better formatting")
        return table
    except Exception as e:
        print(f"Error enhancing table appearance: {str(e)}")
        return table

def apply_table_borders(table, border_size=4, border_color="000000"):
    """
    Apply borders to all cells in a table
    
    Args:
        table: The table to modify
        border_size: Border size (default: 4)
        border_color: Border color in hex (default: black)
    """
    # Define border style
    border_style = {
        "sz": border_size, 
        "val": "single", 
        "color": border_color,
        "space": "0"
    }
    
    # Define a thicker border style for outer edges
    outer_border_style = {
        "sz": border_size + 2,  # Slightly thicker for outer borders
        "val": "single", 
        "color": border_color,
        "space": "0"
    }
    
    # Get table dimensions
    row_count = len(table.rows)
    col_count = len(table.columns) if row_count > 0 else 0
    
    # Apply borders to all cells with special handling for edges
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # Determine which borders should be thicker (outer edges)
            top_style = outer_border_style if i == 0 else border_style
            bottom_style = outer_border_style if i == row_count - 1 else border_style
            start_style = outer_border_style if j == 0 else border_style
            end_style = outer_border_style if j == col_count - 1 else border_style
            
            # Apply borders with appropriate styles
            set_cell_border(
                cell,
                top=top_style,
                bottom=bottom_style,
                start=start_style,
                end=end_style,
                insideH=border_style if i < row_count - 1 else None,
                insideV=border_style if j < col_count - 1 else None
            )
    
    # Enhance the table appearance
    enhance_table_appearance(table)
    
    print("Applied borders to table")
    return table

def copy_table(source_table, target_doc):
    """
    Copy a table from source document to target document.
    
    Args:
        source_table: The table to copy
        target_doc: The document to copy to
    
    Returns:
        The new table in the target document
    """
    # Create a new table with the same dimensions
    new_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    
    # Copy cell contents and formatting
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            # Copy cell text
            target_cell = new_table.cell(i, j)
            if cell.paragraphs:
                for para in cell.paragraphs:
                    # Skip empty paragraphs
                    if not para.text.strip():
                        continue
                    
                    # Create new paragraph in target cell
                    new_para = target_cell.paragraphs[0] if j == 0 and i == 0 else target_cell.add_paragraph()
                    
                    # Copy text with formatting
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        # Copy basic run formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Ensure font color is black
                        new_run.font.color.rgb = RGBColor(0, 0, 0)
                        # Copy font size if available
                        if run.font.size:
                            new_run.font.size = run.font.size
    
    # Apply borders to the table
    apply_table_borders(new_table)
    
    return new_table

def copy_image(paragraph, target_doc):
    """
    Create a new paragraph in target document and copy image from source.
    Uses a more compatible approach to detect and copy images.
    
    Args:
        paragraph: The paragraph containing the image
        target_doc: The document to copy to
    
    Returns:
        bool: True if image was successfully copied
    """
    try:
        # Check if the paragraph has any inline shapes or images
        has_image = False
        image_content = ""
        
        # In python-docx, we can check for images/shapes via runs
        for run in paragraph.runs:
            run_text = run.text.strip()
            # Keep track of any text in the run that contains the image
            if run_text:
                image_content += run_text + " "
            
            # Alternative way to check for images - checking if _element has any children with specific tags
            if hasattr(run, '_element'):
                # Look for common image or shape tags in the element's children
                for child in run._element:
                    tag = child.tag.lower() if hasattr(child, 'tag') else ""
                    # These are common tags for images in docx documents
                    if any(img_tag in tag for img_tag in ['}drawing', '}object', '}picture', '}inline']):
                        has_image = True
                        break
            
            if has_image:
                break
        
        if has_image:
            # Create a new paragraph for the image
            new_para = target_doc.add_paragraph()
            
            # Copy the entire paragraph with its formatting
            # This should maintain the image embedding since we're copying the XML structure
            for run in paragraph.runs:
                new_run = new_para.add_run(run.text)
                # Copy basic formatting
                if hasattr(run, 'bold'):
                    new_run.bold = run.bold
                if hasattr(run, 'italic'):
                    new_run.italic = run.italic
                if hasattr(run, 'underline'):
                    new_run.underline = run.underline
                if hasattr(run, 'font') and hasattr(run.font, 'size'):
                    if run.font.size:
                        new_run.font.size = run.font.size
                new_run.font.color.rgb = RGBColor(0, 0, 0)
            
            print(f"Image detected and preserved in document with content: {image_content.strip()}")
            return True
        
        return False
    except Exception as e:
        print(f"Error in copy_image: {str(e)}")
        return False

def has_image(paragraph):
    """
    Check if a paragraph contains an image using a compatible approach.
    
    Args:
        paragraph: The paragraph to check
        
    Returns:
        bool: True if the paragraph contains an image
    """
    try:
        # Check each run in the paragraph for image content
        for run in paragraph.runs:
            # Check if the run element has any children with image-related tags
            if hasattr(run, '_element'):
                for child in run._element:
                    tag = child.tag.lower() if hasattr(child, 'tag') else ""
                    # These are common tags for images in docx documents
                    if any(img_tag in tag for img_tag in ['}drawing', '}object', '}picture', '}inline']):
                        return True
        return False
    except Exception as e:
        print(f"Error checking for image: {str(e)}")
        return False

def detect_and_create_table_from_text(text, target_doc):
    """
    Attempt to detect tabular data in text and create a table.
    
    Args:
        text (str): Text that might contain tabular data
        target_doc: The document to add the table to
        
    Returns:
        bool: True if a table was created, False otherwise
    """
    lines = text.strip().split('\n')
    if len(lines) < 2:  # Need at least 2 lines for a table
        return False
    
    # Check for common table delimiters
    delimiters = ['\t', '|', ',', '  ', ' {2,}']  # Tab, pipe, comma, double spaces, multiple spaces
    
    for delimiter in delimiters:
        if not delimiter:
            continue
            
        # Split each line by the delimiter
        split_lines = []
        column_counts = []
        
        for line in lines:
            # Handle multiple spaces differently
            if delimiter in [' {2,}', '  ']:
                # Split on multiple spaces (2 or more)
                parts = re.split(r'\s{2,}', line.strip())
            else:
                parts = line.strip().split(delimiter)
            
            # Remove empty parts and extra whitespace
            parts = [p.strip() for p in parts if p.strip()]
            
            if parts:  # Only consider non-empty lines
                split_lines.append(parts)
                column_counts.append(len(parts))
        
        # Check if we have consistent column counts
        if split_lines and len(set(column_counts)) <= 2:  # Allow at most 2 different column counts
            # Get the most common column count
            most_common_count = max(set(column_counts), key=column_counts.count)
            
            # Only create a table if we have at least 2 columns
            if most_common_count >= 2:
                # Create a table with the detected rows and columns
                table = target_doc.add_table(rows=len(split_lines), cols=most_common_count)
                table.style = 'Table Grid'  # Apply a basic table style with borders
                
                # Fill the table with data
                for i, row_data in enumerate(split_lines):
                    for j, cell_text in enumerate(row_data):
                        if j < most_common_count:  # Ensure we don't exceed column count
                            cell = table.cell(i, j)
                            cell.text = cell_text
                
                            # Format cell text
                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            # Make first row bold (header)
                            if i == 0:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(11)
                            else:
                                for run in paragraph.runs:
                                    run.font.size = Pt(11)
                
                # Apply borders and enhance appearance
                apply_table_borders(table)
                enhance_table_appearance(table)
                
                # Add spacing before and after table
                table._element.getparent().addprevious(target_doc.add_paragraph()._element)
                target_doc.add_paragraph()
                
                print(f"Created table with {len(split_lines)} rows and {most_common_count} columns")
                return True
    
    # Additional check for fixed-width column data
    # This helps detect tables that use spaces for alignment
    if all(len(line) > 20 for line in lines):  # Only check longer lines
        # Try to detect columns by looking for aligned spaces
        space_positions = []
        for line in lines:
            positions = [i for i, char in enumerate(line) if char == ' ' and i > 0 and i < len(line)-1]
            space_positions.append(positions)
        
        # Find common space positions
        if space_positions:
            common_spaces = set(space_positions[0])
            for pos_list in space_positions[1:]:
                common_spaces.intersection_update(set(pos_list))
            
            if len(common_spaces) >= 1:  # At least one common space position
                # Split lines at common space positions
                split_lines = []
                positions = sorted(list(common_spaces))
                for line in lines:
                    parts = []
                    start = 0
                    for pos in positions:
                        parts.append(line[start:pos].strip())
                        start = pos
                    parts.append(line[start:].strip())
                    parts = [p for p in parts if p]  # Remove empty parts
                    if parts:
                        split_lines.append(parts)
                
                if split_lines and len(split_lines[0]) >= 2:
                    # Create and format table
                    table = target_doc.add_table(rows=len(split_lines), cols=len(split_lines[0]))
                    table.style = 'Table Grid'
                    
                    # Fill and format table
                    for i, row_data in enumerate(split_lines):
                        for j, cell_text in enumerate(row_data):
                            cell = table.cell(i, j)
                            cell.text = cell_text
                            
                            # Format cell text
                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            # Make first row bold
                            if i == 0:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(11)
                            else:
                                for run in paragraph.runs:
                                    run.font.size = Pt(11)
                    
                    # Apply borders and enhance appearance
                    apply_table_borders(table)
                    enhance_table_appearance(table)
                    
                    # Add spacing
                    table._element.getparent().addprevious(target_doc.add_paragraph()._element)
                    target_doc.add_paragraph()
                    
                    print(f"Created fixed-width table with {len(split_lines)} rows and {len(split_lines[0])} columns")
                return True
    
    return False

def improve_table_formatting(table):
    """
    Improve the formatting of an existing table by ensuring it has proper borders and styling.
    
    Args:
        table: The table to improve
        
    Returns:
        The improved table
    """
    # Check if the table already has a style with borders
    has_borders = False
    
    # Try to determine if the table has borders by checking the first cell
    if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
        first_cell = table.rows[0].cells[0]
        if hasattr(first_cell, '_tc') and hasattr(first_cell._tc, 'tcPr'):
            tcPr = first_cell._tc.tcPr
            if tcPr is not None:
                # Check for existing borders
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                has_borders = tcBorders is not None
    
    # If the table doesn't have borders, apply our border styling
    if not has_borders:
        apply_table_borders(table)
    else:
        # Even if it has borders, enhance its appearance
        enhance_table_appearance(table)
    
    return table

def extract_images_from_pdf(pdf_path, output_folder):
    """
    Extract images from a PDF file using PyMuPDF (fitz).
    
    Args:
        pdf_path (str): Path to the PDF file
        output_folder (str): Folder to save extracted images
        
    Returns:
        list: List of paths to extracted images
    """
    try:
        # Create output folder if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        # Track extracted images to prevent duplicates
        extracted_images = []
        processed_images = set()
        
        # Open the PDF file
        pdf_document = fitz.open(pdf_path)
        
        # Loop through each page
        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            image_list = page.get_images()
            
            # Extract each image
            for image_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    
                    # Skip if we've already processed this image
                    if xref in processed_images:
                        continue
                    processed_images.add(xref)
                    
                    # Get image info
                    base_image = pdf_document.extract_image(xref)
                    if base_image:
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Generate unique filename
                        image_filename = f"page{page_number + 1}_img{image_index + 1}.{image_ext}"
                        image_path = os.path.join(output_folder, image_filename)
                        
                        # Save image
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_bytes)
                        
                        extracted_images.append(image_path)
                        print(f"Extracted image: {image_path}")
                except Exception as img_error:
                    print(f"Error extracting image {image_index + 1} from page {page_number + 1}: {str(img_error)}")
                    continue
        
        pdf_document.close()
        print(f"Total images extracted from PDF: {len(extracted_images)}")
        return extracted_images
    
    except Exception as e:
        print(f"Error extracting images from PDF: {str(e)}")
        return []

def insert_image_into_document(doc, image_path, width=None):
    """
    Insert an image into a Word document with proper formatting.
    
    Args:
        doc: The document to add the image to
        image_path (str): Path to the image file
        width (float, optional): Width in inches for the image. If None, uses 6 inches
        
    Returns:
        bool: True if image was inserted successfully
    """
    try:
        # Add a paragraph for the image
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the image
        run = paragraph.add_run()
        
        # If width not specified, use 6 inches or adjust based on image size
        if width is None:
            width = Inches(6)
        
        # Add the image with specified width
        run.add_picture(image_path, width=width)
        
        # Add some spacing after the image
        paragraph.paragraph_format.space_after = Pt(12)
        return True
    
    except Exception as e:
        print(f"Error inserting image: {str(e)}")
        return False

def extract_images_from_docx(docx_path, output_folder):
    """
    Extract images from a Word document using python-docx and PIL.
    
    Args:
        docx_path (str): Path to the Word document
        output_folder (str): Folder to save extracted images
        
    Returns:
        list: List of paths to extracted images
    """
    try:
        # Create output folder if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        # Load the document
        doc = docx.Document(docx_path)
        extracted_images = []
        processed_refs = set()  # Track processed image references
        
        # Extract images from the document's relationships
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # Skip if we've already processed this image reference
                    if rel.target_ref in processed_refs:
                        continue
                    processed_refs.add(rel.target_ref)
                    
                    # Get image data
                    image_data = rel.target_part.blob
                    image_stream = io.BytesIO(image_data)
                    img = Image.open(image_stream)
                    
                    # Generate filename from the original reference
                    image_filename = os.path.join(output_folder, rel.target_ref.split("/")[-1])
                    base_name, ext = os.path.splitext(image_filename)
                    
                    # Ensure unique filename
                    counter = 1
                    while os.path.exists(image_filename):
                        image_filename = f"{base_name}_{counter}{ext}"
                        counter += 1
                    
                    # Save image
                    img.save(image_filename)
                    extracted_images.append(image_filename)
                    print(f"Extracted image: {image_filename}")
                    
                except Exception as img_error:
                    print(f"Error extracting image {rel.target_ref}: {str(img_error)}")
                    continue
        
        print(f"Total images extracted from DOCX: {len(extracted_images)}")
        return extracted_images
    
    except Exception as e:
        print(f"Error extracting images from Word document: {str(e)}")
        return []

def copy_document_to_template(source_file, template_path="cybergen-template.docx", output_filename="generated_document.docx"):
    try:
        # Check if files exist
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        if not os.path.exists(source_file):
            raise FileNotFoundError(f"Source file not found: {source_file}")
        
        # Load the template document
        template_doc = docx.Document(template_path)
        
        # Set page size to A4 and appropriate margins
        set_page_size_and_margins(template_doc)
        
        # Create output folder for extracted images
        output_folder = os.path.join(os.path.dirname(output_filename), 'extracted_images')
        os.makedirs(output_folder, exist_ok=True)
        
        # Track processed images to prevent duplicates
        processed_images = set()
        
        # Determine file type and process accordingly
        file_ext = os.path.splitext(source_file.lower())[1]
        
        if file_ext == '.pdf':
            print("Processing PDF file...")
            # Extract images from PDF first
            extracted_images = extract_images_from_pdf(source_file, output_folder)
            print(f"Found {len(extracted_images)} images in PDF")
            
            # Convert PDF to DOCX for text and structure
            docx_path = convert_pdf_to_docx(source_file)
            if not docx_path or not os.path.exists(docx_path):
                raise Exception("Failed to convert PDF to DOCX")
            
            # Process the converted DOCX file
            source_doc = docx.Document(docx_path)
            
            # Process document elements in order
            current_image_index = 0
            
            for element in source_doc.element.body:
                if element.tag.endswith('}p'):  # Paragraph
                    # Find corresponding paragraph object
                    para = None
                    for p in source_doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    
                    if not para or not para.text.strip():
                        continue
                    
                    print(f"Processing paragraph: {para.text[:50]}...")
                    
                    # Check if we should insert an image here
                    if current_image_index < len(extracted_images) and has_image(para):
                        img_path = extracted_images[current_image_index]
                        if img_path not in processed_images:
                            print(f"Inserting image: {img_path}")
                            insert_image_into_document(template_doc, img_path)
                            processed_images.add(img_path)
                            current_image_index += 1
                        continue
                    
                    # Process text content
                    new_para = template_doc.add_paragraph()
                    heading_status = is_heading(para.text)
                    
                    if para.runs:
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            if heading_status:
                                new_run.bold = True
                                new_run.font.size = Pt(14)
                            else:
                                new_run.bold = run.bold
                                new_run.font.size = Pt(12.5)
                            new_run.italic = run.italic
                            new_run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        new_run = new_para.add_run(para.text)
                        if heading_status:
                            new_run.bold = True
                            new_run.font.size = Pt(14)
                        else:
                            new_run.bold = False
                            new_run.font.size = Pt(12.5)
                        new_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_space_after_paragraph(new_para, is_heading=heading_status)
                
                elif element.tag.endswith('}tbl'):  # Table
                    # Find corresponding table object
                    table = None
                    for tbl in source_doc.tables:
                        if tbl._element == element:
                            table = tbl
                            break
                    
                    if table:
                        print(f"Processing table with {len(table.rows)} rows and {len(table.columns)} columns")
                        # Add a small space before table
                        template_doc.add_paragraph().paragraph_format.space_after = Pt(6)
                        # Copy and format table
                        new_table = copy_table(table, template_doc)
                        improve_table_formatting(new_table)
                        # Add a small space after table
                        template_doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
            # Insert any remaining images at the end
            while current_image_index < len(extracted_images):
                img_path = extracted_images[current_image_index]
                if img_path not in processed_images:
                    print(f"Inserting remaining image: {img_path}")
                    insert_image_into_document(template_doc, img_path)
                    processed_images.add(img_path)
                current_image_index += 1
            
            # Clean up the temporary converted file
            try:
                os.remove(docx_path)
            except:
                pass
                
        elif file_ext in ('.docx', '.doc'):
            print("Processing Word document...")
            # Extract images from Word document first
            extracted_images = extract_images_from_docx(source_file, output_folder)
            print(f"Found {len(extracted_images)} images in Word document")
            
            # Process Word documents directly
            source_doc = docx.Document(source_file)
            current_image_index = 0
            
            # Process document elements in order
            for element in source_doc.element.body:
                if element.tag.endswith('}p'):  # Paragraph
                    # Find corresponding paragraph object
                    para = None
                    for p in source_doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    
                    if not para or not para.text.strip():
                        continue
                    
                    print(f"Processing paragraph: {para.text[:50]}...")
                    
                    # Check if we should insert an image here
                    if current_image_index < len(extracted_images) and has_image(para):
                        img_path = extracted_images[current_image_index]
                        if img_path not in processed_images:
                            print(f"Inserting image: {img_path}")
                            insert_image_into_document(template_doc, img_path)
                            processed_images.add(img_path)
                            current_image_index += 1
                        continue
                    
                    new_para = template_doc.add_paragraph()
                    heading_status = is_heading(para.text)
                    
                    if para.runs:
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            if heading_status:
                                new_run.bold = True
                                new_run.font.size = Pt(14)
                            else:
                                new_run.bold = run.bold
                                new_run.font.size = Pt(12.5)
                            new_run.italic = run.italic
                            new_run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        new_run = new_para.add_run(para.text)
                        if heading_status:
                            new_run.bold = True
                            new_run.font.size = Pt(14)
                        else:
                            new_run.bold = False
                            new_run.font.size = Pt(12.5)
                        new_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_space_after_paragraph(new_para, is_heading=heading_status)
                
                elif element.tag.endswith('}tbl'):  # Table
                    # Find corresponding table object
                    table = None
                    for tbl in source_doc.tables:
                        if tbl._element == element:
                            table = tbl
                            break
                    
                    if table:
                        print(f"Processing table with {len(table.rows)} rows and {len(table.columns)} columns")
                        # Add a small space before table
                        template_doc.add_paragraph().paragraph_format.space_after = Pt(6)
                        # Copy and format table
                        new_table = copy_table(table, template_doc)
                        improve_table_formatting(new_table)
                        # Add a small space after table
                        template_doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
            # Insert any remaining images at the end
            while current_image_index < len(extracted_images):
                img_path = extracted_images[current_image_index]
                if img_path not in processed_images:
                    print(f"Inserting remaining image: {img_path}")
                    insert_image_into_document(template_doc, img_path)
                    processed_images.add(img_path)
                current_image_index += 1
        
        # Set widow/orphan control
        for paragraph in template_doc.paragraphs:
            paragraph.paragraph_format.widow_control = True
        
        # Save the document
        template_doc.save(output_filename)
        print(f"Document saved to: {output_filename}")
        return os.path.abspath(output_filename)
    
    except Exception as e:
        print(f"Error copying document: {str(e)}")
        return None

def insert_text_into_template(input_text, template_path="cybergen-template.docx", output_filename="generated_document.docx"):
    """
    Inserts the user's text into the template document.
    
    Args:
        input_text (str): The text content to be inserted
        template_path (str): Path to the template document
        output_filename (str): Name for the output document
    
    Returns:
        str: Path to the created document
    """
    try:
        # Check if template exists
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        print("Input text for document creation:")
        print(input_text[:500] + "..." if len(input_text) > 500 else input_text)
        
        # Load the template document
        doc = docx.Document(template_path)
        
        # Set page size to A4 and appropriate margins
        set_page_size_and_margins(doc)
        
        # Process the input text
        paragraphs = input_text.strip().split('\n\n')  # Split on double newlines to better identify paragraphs
        
        # Keep track of the last few paragraphs to help identify subheadings
        recent_paragraphs = []
        heading_count = 0  # Track how many headings we've seen recently
        
        # Process each paragraph
        i = 0
        while i < len(paragraphs):
            paragraph_text = paragraphs[i]
            if not paragraph_text.strip():  # Skip empty paragraphs
                i += 1
                continue
            
            print(f"Processing paragraph {i}: {paragraph_text[:50]}...")
            
            # Try to detect and create a table from this paragraph
            # If it looks like tabular data with multiple lines
            if '\n' in paragraph_text and len(paragraph_text.split('\n')) > 1:
                if detect_and_create_table_from_text(paragraph_text, doc):
                    # Table was created, move to next paragraph
                    i += 1
                    continue
            
            # Check if this paragraph is a heading or subheading
            heading_status = is_heading(paragraph_text)
            
            # Check for subheading based on context (position relative to headings)
            subheading_status = False
            if heading_status:
                heading_count = 3  # Reset counter for future subheadings
                print(f"Heading detected: {paragraph_text}")
            elif heading_count > 0:
                subheading_status = is_subheading(paragraph_text, recent_paragraphs, True)
                if subheading_status:
                    print(f"Subheading detected (after heading): {paragraph_text}")
                heading_count -= 1
            else:
                subheading_status = is_subheading(paragraph_text, recent_paragraphs, False)
                if subheading_status:
                    print(f"Subheading detected (standalone): {paragraph_text}")
            
            # Keep track of recent paragraphs for context
            recent_paragraphs.append(paragraph_text)
            if len(recent_paragraphs) > 5:
                recent_paragraphs.pop(0)
            
            # Add paragraph with appropriate formatting
            p = doc.add_paragraph()
            if len(doc.paragraphs) > 1:
                # Copy style from an existing paragraph if available
                p.style = doc.paragraphs[1].style
            
            # Add run with appropriate formatting
            run = p.add_run(paragraph_text)
            
            # Apply formatting based on heading/subheading status
            if heading_status:
                run.font.size = Pt(14)
                run.bold = True
                run.underline = WD_UNDERLINE.SINGLE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif subheading_status:
                run.font.size = Pt(13)  # 1pt smaller than headings
                run.bold = True
                run.underline = WD_UNDERLINE.SINGLE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run.font.size = Pt(12.5)
                run.bold = False
                run.underline = None
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add proper spacing after paragraph using Word's standard
            add_space_after_paragraph(p, is_heading=heading_status, is_subheading=subheading_status)
            
            i += 1
        
        # Set widow/orphan control for the whole document to prevent single lines
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.widow_control = True
        
        # Save the document
        doc.save(output_filename)
        print(f"Document saved to: {output_filename}")
        return os.path.abspath(output_filename)
    
    except Exception as e:
        print(f"Error creating document: {str(e)}")
        return None

def main():
    """
    Main function to handle document processing.
    """
    print("CyberGen Document Formatter")
    print("==========================")
    
    # Default template path
    template_path = "cybergen-template.docx"
    
    while True:
        print("\nOptions:")
        print("1. Import text from a document (Word or PDF)")
        print("2. Exit")
        
        choice = input("\nEnter your choice (1-2): ")
        
        if choice == '1':
            file_path = input("\nEnter the path to the document (Word or PDF): ")
            
            if not os.path.exists(file_path):
                print(f"Error: File not found at {file_path}")
                continue
                
            output_name = input("\nEnter output filename (leave blank for default 'generated_document.docx'): ")
            if not output_name:
                output_name = "generated_document.docx"
            elif not output_name.lower().endswith('.docx'):
                output_name += '.docx'
            
            # Use the function to copy content preserving formatting
            document_path = copy_document_to_template(file_path, template_path=template_path, output_filename=output_name)
            if document_path:
                print(f"\nDocument successfully created at: {document_path}")
                print("Note: Text has been formatted according to heading detection rules.")
                print("      All paragraphs have standard spacing after them.")
                print("      Headings are kept with their following paragraphs across page breaks.")
        
        elif choice == '2':
            print("\nExiting program. Goodbye!")
            break
        
        else:
            print("\nInvalid choice. Please try again.")

if __name__ == "__main__":
    main() 