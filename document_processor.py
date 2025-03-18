from enum import Enum
import shutil
from pathlib import Path
import asyncio
from typing import Dict, Optional
import os
import docx
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pdf2docx import Converter
import fitz
from PIL import Image
import io
from datetime import datetime

# Import all the functions from cybergen_template
from cybergen_template import (
    set_page_size_and_margins,
    is_heading,
    is_subheading,
    format_paragraph,
    add_space_after_paragraph,
    enhance_table_appearance,
    apply_table_borders,
    copy_table,
    copy_image,
    has_image,
    detect_and_create_table_from_text,
    extract_images_from_pdf,
    extract_images_from_docx,
    insert_image_into_document
)

class JobStatus(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"

def get_job_status(job_id: str, jobs: Dict) -> Optional[Dict]:
    """
    Get the status of a job.
    """
    if job_id not in jobs:
        return None
    return jobs[job_id]

async def cleanup_job(job_id: str, jobs: Dict, upload_dir: Path, output_dir: Path):
    """
    Clean up job files and data.
    """
    # Remove job directories
    job_upload_dir = upload_dir / job_id
    job_output_dir = output_dir / job_id
    
    if job_upload_dir.exists():
        shutil.rmtree(job_upload_dir)
    if job_output_dir.exists():
        shutil.rmtree(job_output_dir)
    
    # Remove job from jobs dict
    if job_id in jobs:
        del jobs[job_id]

def convert_pdf_to_docx(pdf_path: Path, output_path: Path) -> Optional[Path]:
    """
    Convert PDF to DOCX format.
    """
    try:
        cv = Converter(str(pdf_path))
        cv.convert(str(output_path))
        cv.close()
        return output_path
    except Exception as e:
        print(f"Error converting PDF to DOCX: {str(e)}")
        return None

def add_page_number(paragraph):
    """Add page numbers in the footer."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def setup_headers_and_footers(doc, title="", author="", date=None):
    """Setup headers and footers for the document."""
    # Set up header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add title to header
    if title:
        title_run = header_para.add_run(title)
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Set up footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add author and date to left side of footer
    if author or date:
        left_footer = f"{author}"
        if date:
            left_footer += f" - {date}"
        if left_footer:
            footer_para.add_run(left_footer).font.size = Pt(10)
    
    # Add page numbers to right side of footer
    footer_para.add_run("Page ").font.size = Pt(10)
    add_page_number(footer_para)

def process_document(job_id: str, input_dir: Path, output_dir: Path, jobs: Dict):
    """
    Process all uploaded documents in the input directory.
    """
    try:
        # Update job status
        jobs[job_id]["status"] = JobStatus.PROCESSING
        print(f"\nStarting job {job_id}")
        
        # Find the template file
        template_path = Path(__file__).parent / "cybergen-template.docx"
        if not template_path.exists():
            # Try current directory
            template_path = Path("cybergen-template.docx")
            if not template_path.exists():
                raise Exception(f"Template file not found in {template_path} or current directory")
        
        print(f"Using template from: {template_path}")
        
        # Get all input files
        input_files = list(input_dir.glob("*.*"))
        if not input_files:
            raise Exception("No input files found in upload directory")
        
        print(f"Found {len(input_files)} files to process")
        
        processed_files = []
        error_files = []
        
        # Process each input file
        for input_file in input_files:
            try:
                print(f"\nProcessing input file: {input_file}")
                print(f"File size: {input_file.stat().st_size} bytes")
                print(f"File permissions: {oct(input_file.stat().st_mode)}")
                
                # Validate file
                if not input_file.exists():
                    raise Exception("File does not exist")
                if not input_file.is_file():
                    raise Exception("Not a valid file")
                if input_file.stat().st_size == 0:
                    raise Exception("File is empty")
                if not input_file.suffix.lower() in ['.docx', '.pdf']:
                    raise Exception(f"Unsupported file type: {input_file.suffix}")
                
                # Create a copy of the input file with proper permissions
                temp_input_file = input_dir / f"temp_{input_file.name}"
                try:
                    shutil.copy2(input_file, temp_input_file)
                    os.chmod(temp_input_file, 0o644)  # Set read/write permissions
                    print(f"Created temporary file with proper permissions: {temp_input_file}")
                    print(f"Temp file size: {temp_input_file.stat().st_size} bytes")
                    print(f"Temp file permissions: {oct(temp_input_file.stat().st_mode)}")
                except Exception as e:
                    raise Exception(f"Failed to create temporary file with proper permissions: {str(e)}")
                
                # Create template document from existing template
                try:
                    template_doc = docx.Document(template_path)
                    print("Successfully loaded template document")
                except Exception as e:
                    if temp_input_file.exists():
                        temp_input_file.unlink()  # Clean up temp file
                    raise Exception(f"Failed to load template document: {str(e)}")
                
                # Clear template content while preserving headers/footers
                try:
                    for paragraph in template_doc.paragraphs[:-1]:  # Keep last paragraph
                        p = paragraph._element
                        p.getparent().remove(p)
                    print("Successfully cleared template content")
                except Exception as e:
                    print(f"Warning: Error clearing template content: {str(e)}")
                
                # Process based on file type
                # Create images directory for this file
                images_dir = output_dir / "images" / input_file.stem
                images_dir.mkdir(exist_ok=True, parents=True)
                extracted_images = []

                try:
                    if temp_input_file.suffix.lower() == '.pdf':
                        print("Processing PDF file...")
                        # Extract images from PDF
                        try:
                            extracted_images = extract_images_from_pdf(str(temp_input_file), str(images_dir))
                            print(f"Extracted {len(extracted_images)} images from PDF")
                        except Exception as e:
                            print(f"Warning: Failed to extract images from PDF: {str(e)}")
                            extracted_images = []
                        
                        # Convert PDF to DOCX
                        docx_path = output_dir / f"{temp_input_file.stem}_converted.docx"
                        if not convert_pdf_to_docx(temp_input_file, docx_path):
                            raise Exception("Failed to convert PDF to DOCX")
                        print("Successfully converted PDF to DOCX")
                        
                        try:
                            source_doc = docx.Document(docx_path)
                            print("Successfully loaded converted DOCX")
                        except Exception as e:
                            raise Exception(f"Failed to open converted DOCX: {str(e)}")
                    else:
                        print("Processing DOCX file...")
                        try:
                            source_doc = docx.Document(temp_input_file)
                            print("Successfully loaded source DOCX")
                            print(f"Document sections: {len(source_doc.sections)}")
                            print(f"Document paragraphs: {len(source_doc.paragraphs)}")
                            print(f"Document tables: {len(source_doc.tables)}")
                        except Exception as e:
                            print(f"Error details for DOCX: {type(e).__name__}: {str(e)}")
                            raise Exception(f"Failed to open source DOCX: {str(e)}")
                        
                        # Extract images after successfully loading document
                        try:
                            extracted_images = extract_images_from_docx(str(temp_input_file), str(images_dir))
                            print(f"Extracted {len(extracted_images)} images from DOCX")
                        except Exception as e:
                            print(f"Warning: Failed to extract images from DOCX: {str(e)}")
                            extracted_images = []
                except Exception as e:
                    raise Exception(f"Failed to process document: {str(e)}")

                # Check if Table Grid style exists in template
                table_grid_style_exists = 'Table Grid' in template_doc.styles
                print(f"Table Grid style exists: {table_grid_style_exists}")

                # Track images and headings
                current_image_index = 0
                recent_paragraphs = []
                heading_count = 0
                
                # Process document content
                for element in source_doc.element.body:
                    try:
                        if element.tag.endswith('}p'):  # Paragraph
                            # Find corresponding paragraph
                            para = None
                            for p in source_doc.paragraphs:
                                if p._element == element:
                                    para = p
                                    break
                            
                            if not para or not para.text.strip():
                                continue
                            
                            # Check for images
                            if current_image_index < len(extracted_images) and has_image(para):
                                try:
                                    img_path = extracted_images[current_image_index]
                                    insert_image_into_document(template_doc, img_path)
                                    current_image_index += 1
                                    continue
                                except Exception as e:
                                    print(f"Warning: Failed to insert image: {str(e)}")
                            
                            # Check for tables in text
                            if '\n' in para.text and len(para.text.split('\n')) > 1:
                                try:
                                    if detect_and_create_table_from_text(para.text, template_doc):
                                        continue
                                except Exception as e:
                                    print(f"Warning: Failed to create table from text: {str(e)}")
                            
                            # Process paragraph
                            heading_status = is_heading(para.text)
                            subheading_status = False
                            
                            if heading_status:
                                heading_count = 3
                            elif heading_count > 0:
                                subheading_status = is_subheading(para.text, recent_paragraphs, True)
                                heading_count -= 1
                            else:
                                subheading_status = is_subheading(para.text, recent_paragraphs, False)
                            
                            # Keep track of recent paragraphs
                            recent_paragraphs.append(para.text)
                            if len(recent_paragraphs) > 5:
                                recent_paragraphs.pop(0)
                            
                            # Add formatted paragraph
                            try:
                                new_para = template_doc.add_paragraph()
                                format_paragraph(new_para, heading_status, subheading_status)
                                
                                # Copy text with formatting
                                for run in para.runs:
                                    new_run = new_para.add_run(run.text)
                                    if heading_status:
                                        new_run.font.size = Pt(14)
                                        new_run.bold = True
                                    elif subheading_status:
                                        new_run.font.size = Pt(13)
                                        new_run.bold = True
                                    else:
                                        new_run.font.size = Pt(12)
                                        new_run.bold = run.bold
                                    new_run.italic = run.italic
                                    new_run.font.color.rgb = RGBColor(0, 0, 0)
                                
                                add_space_after_paragraph(new_para, heading_status, subheading_status)
                            except Exception as e:
                                print(f"Warning: Failed to format paragraph: {str(e)}")
                        
                        elif element.tag.endswith('}tbl'):  # Table
                            # Find corresponding table
                            table = None
                            for tbl in source_doc.tables:
                                if tbl._element == element:
                                    table = tbl
                                    break
                            
                            if table:
                                try:
                                    # Add spacing before table
                                    template_doc.add_paragraph().paragraph_format.space_after = Pt(6)
                                    # Copy and format table
                                    new_table = copy_table(table, template_doc)
                                    
                                    # Apply table style if available
                                    if table_grid_style_exists:
                                        try:
                                            new_table.style = 'Table Grid'
                                        except Exception as e:
                                            print(f"Warning: Failed to apply Table Grid style: {str(e)}")
                                    
                                    enhance_table_appearance(new_table)
                                    apply_table_borders(new_table)
                                    # Add spacing after table
                                    template_doc.add_paragraph().paragraph_format.space_after = Pt(6)
                                except Exception as e:
                                    print(f"Warning: Failed to process table: {str(e)}")
                    except Exception as e:
                        print(f"Warning: Failed to process document element: {str(e)}")
                        continue
                
                # Insert any remaining images
                while current_image_index < len(extracted_images):
                    try:
                        img_path = extracted_images[current_image_index]
                        insert_image_into_document(template_doc, img_path)
                        current_image_index += 1
                    except Exception as e:
                        print(f"Warning: Failed to insert remaining image: {str(e)}")
                        current_image_index += 1
                
                # Save processed document
                output_file = output_dir / f"processed_{input_file.name}"
                if output_file.suffix.lower() != '.docx':
                    output_file = output_file.with_suffix('.docx')
                
                try:
                    # Ensure output directory has proper permissions
                    output_dir.chmod(0o755)  # Directory needs execute permission
                    template_doc.save(output_file)
                    output_file.chmod(0o644)  # Make output file readable
                    print(f"Successfully saved processed document to {output_file}")
                    print(f"Output file size: {output_file.stat().st_size} bytes")
                    processed_files.append(output_file.name)
                except Exception as e:
                    print(f"Error saving output file: {type(e).__name__}: {str(e)}")
                    error_files.append({
                        "filename": input_file.name,
                        "error": f"Failed to save output document: {str(e)}"
                    })
                    continue
                
            except Exception as file_error:
                error_msg = f"{type(file_error).__name__}: {str(file_error)}"
                print(f"Error processing file {input_file.name}: {error_msg}")
                error_files.append({
                    "filename": input_file.name,
                    "error": error_msg
                })
                continue
            finally:
                # Clean up temporary files
                if 'temp_input_file' in locals() and temp_input_file.exists():
                    try:
                        temp_input_file.unlink()
                        print("Cleaned up temporary input file")
                    except Exception as e:
                        print(f"Warning: Failed to clean up temporary file: {str(e)}")
                if 'docx_path' in locals() and docx_path.exists():
                    try:
                        docx_path.unlink()
                        print("Cleaned up converted DOCX file")
                    except Exception as e:
                        print(f"Warning: Failed to clean up converted file: {str(e)}")
        
        # Check if any files were processed successfully
        if not processed_files:
            raise Exception("No files were processed successfully")
        
        # Update job information
        jobs[job_id].update({
            "status": JobStatus.COMPLETED,
            "output_files": processed_files,
            "error_files": error_files if error_files else None,
            "completed_at": datetime.utcnow().isoformat()
        })
        print(f"Job {job_id} completed successfully. Processed {len(processed_files)} files.")
        
    except Exception as e:
        error_msg = f"{type(e).__name__}: {str(e)}"
        print(f"Error processing documents: {error_msg}")
        # Update job with error information
        jobs[job_id].update({
            "status": JobStatus.FAILED,
            "error": error_msg,
            "error_files": error_files if error_files else None,
            "failed_at": datetime.utcnow().isoformat()
        }) 